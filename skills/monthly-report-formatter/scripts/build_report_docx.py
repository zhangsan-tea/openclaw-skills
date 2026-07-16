#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
月报排版生成器 —— 依据结构化 JSON 生成 Word(.docx)月报。
排版风格参照"阅览版"：蓝底一级章节条、蓝竖线小标题、专业表格、
"值得关注"浅蓝提示框、配图占位、落款等。

用法:
    python build_report_docx.py <input.json> <output.docx>

JSON 结构见 references/sample_report_2026-06.json，节点 type 支持:
    h1 / h2 / h3 / para / callout / note / list / table / figure
"""
import sys
import os
import json
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- 配色（参照阅览版） ----------
C_BLUE_DARK = RGBColor(0x1F, 0x4E, 0x8C)   # 深蓝
C_BLUE = RGBColor(0x2E, 0x5C, 0xA8)        # 小标题蓝
BG_CALLOUT = "EAF1FB"                       # 值得关注框底色
BG_THEAD = "1F4E8C"                         # 表头底色
BG_STRIPE = "F5F8FC"                        # 斑马纹
C_TEXT = RGBColor(0x33, 0x33, 0x33)
C_GRAY = RGBColor(0x88, 0x88, 0x88)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)

FONT_CN = "微软雅黑"
FONT_CN_FALLBACK = "Microsoft YaHei"


def set_cn_font(run, name=FONT_CN):
    """设置中文字体（东亚字体需单独设置）。"""
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), name)
    rfonts.set(qn('w:ascii'), name)
    rfonts.set(qn('w:hAnsi'), name)


def shade_cell(cell, hex_color):
    """给表格单元格填充底色。"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def shade_paragraph(paragraph, hex_color):
    """给段落加底纹（用于 callout 框）。"""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)


def add_left_border(paragraph, hex_color="1F4E8C", size=24):
    """给段落加左边框（模拟竖线小标题 / callout 左边条）。"""
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), str(size))
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), hex_color)
    pbdr.append(left)
    pPr.append(pbdr)


def set_paragraph_spacing(paragraph, before=6, after=6, line=1.5):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


# ---------- 各类节点渲染 ----------

def render_title(doc, meta):
    # 主标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=4)
    run = p.add_run(meta.get("title", ""))
    set_cn_font(run)
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = C_BLUE_DARK
    # 汇报周期
    if meta.get("period"):
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p2, before=0, after=12)
        r2 = p2.add_run(meta["period"])
        set_cn_font(r2)
        r2.font.size = Pt(10.5)
        r2.font.color.rgb = C_GRAY


def render_intro(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=12, line=1.6)
    shade_paragraph(p, BG_CALLOUT)
    p.paragraph_format.left_indent = Pt(6)
    p.paragraph_format.right_indent = Pt(6)
    _add_rich_runs(p, text, base_size=11)


def render_h1(doc, text):
    """一级章节：蓝底白字条。"""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    shade_cell(cell, BG_THEAD)
    # 控制宽度不占满整行
    _set_table_width(table, 8.0)
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, before=3, after=3)
    run = p.add_run("  " + text + "  ")
    set_cn_font(run)
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = C_WHITE
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def render_h2(doc, text):
    """二级标题：蓝竖线 + 蓝字。"""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=8, after=4)
    add_left_border(p, "2E5CA8", size=28)
    p.paragraph_format.left_indent = Pt(8)
    run = p.add_run(text)
    set_cn_font(run)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = C_BLUE


def render_h3(doc, text):
    """三级标题：加粗蓝字。"""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=3)
    run = p.add_run(text)
    set_cn_font(run)
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = C_BLUE_DARK


def render_para(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=3, after=3, line=1.6)
    p.paragraph_format.first_line_indent = Pt(22)
    _add_rich_runs(p, text, base_size=11)


def render_callout(doc, text):
    """值得关注：浅蓝底 + 左蓝条。"""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=4, after=8, line=1.6)
    shade_paragraph(p, BG_CALLOUT)
    add_left_border(p, "2E5CA8", size=36)
    p.paragraph_format.left_indent = Pt(10)
    p.paragraph_format.right_indent = Pt(6)
    _add_rich_runs(p, text, base_size=10.5, callout=True)


def render_note(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=4)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    set_cn_font(run)
    run.font.size = Pt(9.5)
    run.font.italic = True
    run.font.color.rgb = C_GRAY


def render_list(doc, items):
    for it in items:
        p = doc.add_paragraph(style=None)
        set_paragraph_spacing(p, before=2, after=2, line=1.5)
        p.paragraph_format.left_indent = Pt(22)
        p.paragraph_format.first_line_indent = Pt(-10)
        run = p.add_run("● ")
        set_cn_font(run)
        run.font.size = Pt(9)
        run.font.color.rgb = C_BLUE
        _add_rich_runs(p, it, base_size=11)


def render_table(doc, node):
    caption = node.get("caption")
    if caption:
        pc = doc.add_paragraph()
        set_paragraph_spacing(pc, before=6, after=3)
        rc = pc.add_run(caption)
        set_cn_font(rc)
        rc.font.size = Pt(10.5)
        rc.font.bold = True
        rc.font.color.rgb = C_BLUE_DARK

    headers = node["headers"]
    rows = node["rows"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        shade_cell(hdr[i], BG_THEAD)
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_cn_font(run)
        run.font.size = Pt(9.5)
        run.font.bold = True
        run.font.color.rgb = C_WHITE

    # 数据行
    for ri, row in enumerate(rows):
        cells = table.add_row().cells
        for ci, val in enumerate(row):
            if ri % 2 == 1:
                shade_cell(cells[ci], BG_STRIPE)
            p = cells[ci].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci == 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            set_cn_font(run)
            run.font.size = Pt(9)
            run.font.color.rgb = C_TEXT
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def render_figure(doc, node):
    """配图：如有 image 字段则插入真实图片，否则显示占位框。"""
    image_path = node.get("image", "")

    if image_path and os.path.exists(image_path):
        # --- 嵌入真实图片 ---
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, before=8, after=2)

        run = p.add_run()
        try:
            # 最大宽度约 15cm（A4 页面宽度减去边距），自适应比例
            run.add_picture(image_path, width=Cm(15.0))
        except Exception as e:
            # 降级为占位提示
            run_text = p.add_run(f"［图片加载失败: {os.path.basename(image_path)}］")
            set_cn_font(run_text)
            run_text.font.size = Pt(9.5)
            run_text.font.color.rgb = C_GRAY

        # 图注
        if node.get("caption"):
            pc = doc.add_paragraph()
            pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(pc, before=2, after=6)
            rc = pc.add_run(node["caption"])
            set_cn_font(rc)
            rc.font.size = Pt(9.5)
            rc.font.color.rgb = C_GRAY
    else:
        # --- 无图片：占位框 ---
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, before=6, after=2)
        shade_paragraph(p, "F0F3F8")
        r = p.add_run("［" + node.get("note", "配图") + "］")
        set_cn_font(r)
        r.font.size = Pt(9.5)
        r.font.color.rgb = C_GRAY
        if node.get("caption"):
            pc = doc.add_paragraph()
            pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(pc, before=0, after=8)
            rc = pc.add_run(node["caption"])
            set_cn_font(rc)
            rc.font.size = Pt(9.5)
            rc.font.color.rgb = C_GRAY


def render_footer(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=16, after=0)
    r = p.add_run(text)
    set_cn_font(r)
    r.font.size = Pt(10)
    r.font.color.rgb = C_GRAY


# ---------- 富文本：**加粗** 与 "引号" 处理 ----------
import re

def _add_rich_runs(paragraph, text, base_size=11, callout=False):
    """支持 **加粗** 标记；callout 中的加粗用深蓝。"""
    bold_color = C_BLUE_DARK if callout else C_TEXT
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
            run.font.color.rgb = bold_color
        else:
            run = paragraph.add_run(part)
            run.font.color.rgb = C_TEXT
        set_cn_font(run)
        run.font.size = Pt(base_size)


def _set_table_width(table, width_cm):
    for row in table.rows:
        for cell in row.cells:
            cell.width = Cm(width_cm)


# ---------- 遍历渲染 ----------

def render_node(doc, node):
    t = node.get("type")
    if t == "h1":
        render_h1(doc, node["text"])
    elif t == "h2":
        render_h2(doc, node["text"])
    elif t == "h3":
        render_h3(doc, node["text"])
    elif t == "para":
        render_para(doc, node["text"])
    elif t == "callout":
        render_callout(doc, node["text"])
    elif t == "note":
        render_note(doc, node["text"])
    elif t == "list":
        render_list(doc, node["items"])
    elif t == "table":
        render_table(doc, node)
    elif t == "figure":
        render_figure(doc, node)
    for child in node.get("children", []):
        render_node(doc, child)


def build(data, out_path):
    doc = Document()
    # 默认字体
    style = doc.styles['Normal']
    style.font.name = FONT_CN
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CN)
    # 页边距
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    render_title(doc, data["meta"])
    if data.get("intro"):
        intro_text = data["intro"] if isinstance(data["intro"], str) else data["intro"].get("text", "")
        if intro_text:
            render_intro(doc, intro_text)
    for section in data["sections"]:
        render_node(doc, section)
    if data["meta"].get("footer"):
        render_footer(doc, data["meta"]["footer"])

    doc.save(out_path)
    print("已生成 Word:", out_path)


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("用法: python build_report_docx.py <input.json> <output.docx>")
        sys.exit(1)
    with open(sys.argv[1], "r", encoding="utf-8") as f:
        data = json.load(f)
    build(data, sys.argv[2])
